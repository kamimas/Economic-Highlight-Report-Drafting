# -*- coding: utf-8 -*-

# Form implementation generated from reading ui file 'C:\Users\khossei4\Desktop\Ui Files\Electricity.ui'
#
# Created by: PyQt5 UI code generator 5.12.2
#
# WARNING! All changes made in this file will be lost!

from PyQt5 import QtCore, QtGui, QtWidgets
import pandas               as pd  
import matplotlib.pyplot    as plt     
import numpy                as np 
import urllib.request
import os.path
from datetime import date,datetime
from docx import Document
        
class Ui_ElectricityWindow(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(800, 600)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setGeometry(QtCore.QRect(40, 20, 331, 81))
        font = QtGui.QFont()
        font.setPointSize(16)
        self.label.setFont(font)
        self.label.setObjectName("label")
        self.tableWidget = QtWidgets.QTableWidget(self.centralwidget)
        self.tableWidget.setGeometry(QtCore.QRect(30, 110, 561, 151))
        self.tableWidget.setObjectName("tableWidget")
        self.tableWidget.setColumnCount(3)
        self.tableWidget.setRowCount(3)
        item = QtWidgets.QTableWidgetItem()
        self.tableWidget.setVerticalHeaderItem(0, item)
        item = QtWidgets.QTableWidgetItem()
        self.tableWidget.setVerticalHeaderItem(1, item)
        item = QtWidgets.QTableWidgetItem()
        self.tableWidget.setVerticalHeaderItem(2, item)
    
        item = QtWidgets.QTableWidgetItem()
        self.tableWidget.setHorizontalHeaderItem(0, item)
        item = QtWidgets.QTableWidgetItem()
        self.tableWidget.setHorizontalHeaderItem(1, item)
        item = QtWidgets.QTableWidgetItem()
        self.tableWidget.setHorizontalHeaderItem(2, item)
        self.label_3 = QtWidgets.QLabel(self.centralwidget)
        self.label_3.setGeometry(QtCore.QRect(30, 330, 191, 31))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.label_3.setFont(font)
        self.label_3.setObjectName("label_3")
        self.summaryField = QtWidgets.QTextEdit(self.centralwidget)
        self.summaryField.setGeometry(QtCore.QRect(30, 390, 561, 151))
        self.summaryField.setObjectName("summaryField")
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 800, 26))
        self.menubar.setObjectName("menubar")
        self.menuElectricity_Details = QtWidgets.QMenu(self.menubar)
        self.menuElectricity_Details.setObjectName("menuElectricity_Details")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)
        self.menubar.addAction(self.menuElectricity_Details.menuAction())

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)


        self.populate()
    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
        self.label.setText(_translate("MainWindow", "Last Recorded Date"))
        item = self.tableWidget.verticalHeaderItem(0)
        item.setText(_translate("MainWindow", "Peak Demand  (MWh)"))
        item = self.tableWidget.verticalHeaderItem(1)
        item.setText(_translate("MainWindow", "Total Demand (GWh)"))
        item = self.tableWidget.verticalHeaderItem(2)
        item.setText(_translate("MainWindow", "Average Electricity Price"))
    
        item = self.tableWidget.horizontalHeaderItem(0)
        item.setText(_translate("MainWindow", "First Month"))

        item = self.tableWidget.horizontalHeaderItem(1)
        item.setText(_translate("MainWindow", "Second Month"))

        item = self.tableWidget.horizontalHeaderItem(2)
        item.setText(_translate("MainWindow", "Change"))

        self.label_3.setText(_translate("MainWindow", "Summary"))
        self.menuElectricity_Details.setTitle(_translate("MainWindow", "Electricity Details"))


    def populate(self):
        path="HistoricalPoolPriceReportServlet.xlsx"
        xl = pd.ExcelFile(path)

        df1 = xl.parse('Data') 
        LastRecordedDate = df1['Date (HE)'][df1.shape[0]-1]
        LastRecordedDate = datetime.strptime(LastRecordedDate[0:10],'%m/%d/%Y')
        
        self.label.setText(LastRecordedDate.strftime('%B %d, %Y'))


        firstMonthList = []
        secondMonthList = []

        Average = 0
        AverageLastMonth = 0

        thismonth = int(date.today().strftime('%m'))
        lastmonth = int(date.today().strftime('%m'))-1

        # this data is recorded in a list to be sent to the populateTable()
        firstMonthList.append(datetime.strptime(str(thismonth),'%m').strftime('%B'))
        secondMonthList.append(datetime.strptime(str(lastmonth),'%m').strftime('%B'))

        price =0
        j = 0
        k = 0


        PeakDemandFirstMonth = 0
        TotalDemandFirstMonth = 0


        PeakDemandSecondMonth = 0
        TotalDemandSecondMonth = 0
        # print(int(df1['AIL Demand (MW)'][0]))
        for i in range(1,df1.shape[0]):
            
            startDate = df1['Date (HE)'][i]
            startDate = int(startDate[0:2])
    

            if thismonth == startDate:

                Max = int(df1['AIL Demand (MW)'][i])
                if  Max >= PeakDemandFirstMonth:
                    PeakDemandFirstMonth = Max


                
                TotalDemandFirstMonth += int(df1['AIL Demand (MW)'][i])

                price = df1['Price ($)'][i]
                j+=1
                Average += price
            elif lastmonth == startDate:
                Max = int(df1['AIL Demand (MW)'][i])
                if Max >= PeakDemandSecondMonth:

                    PeakDemandSecondMonth = Max

                TotalDemandSecondMonth += int(df1['AIL Demand (MW)'][i])

                price = df1['Price ($)'][i]
                k+=1
                AverageLastMonth += price
            # startDate = datetime.strptime(startDate, '%m/%d/%Y')


        Average = Average/j
        AverageLastMonth = AverageLastMonth/k
        
        # Average is appended to the list
        firstMonthList.append(Average)
        secondMonthList.append(AverageLastMonth)
        
        firstMonthList.append(TotalDemandFirstMonth)
        secondMonthList.append(TotalDemandSecondMonth)

        firstMonthList.append(PeakDemandFirstMonth)
        secondMonthList.append(PeakDemandSecondMonth)

        self.tableWidget.setHorizontalHeaderItem(1,QtWidgets.QTableWidgetItem(firstMonthList[0] + " " + LastRecordedDate.strftime("%Y") ))
        self.tableWidget.setHorizontalHeaderItem(0,QtWidgets.QTableWidgetItem(secondMonthList[0] + " "+ LastRecordedDate.strftime("%Y"))) 

        self.populateTable(firstMonthList,secondMonthList)

    def populateTable(self,firstMonthList,secondMonthList):

        self.tableWidget.setItem(0,0, QtWidgets.QTableWidgetItem(str(secondMonthList[3])))
        self.tableWidget.setItem(0,1, QtWidgets.QTableWidgetItem(str(firstMonthList[3])))
        change = str(round(((firstMonthList[3]/secondMonthList[3])-1)*100,1))
        self.tableWidget.setItem(0,2, QtWidgets.QTableWidgetItem(change+"%"))


        self.tableWidget.setItem(1,0, QtWidgets.QTableWidgetItem(str(round(secondMonthList[2]/1000,1))))
        self.tableWidget.setItem(1,1, QtWidgets.QTableWidgetItem(str(round(firstMonthList[2]/1000,1))))
        change = str(round(((firstMonthList[2]/secondMonthList[2])-1)*100,1))
        self.tableWidget.setItem(1,2, QtWidgets.QTableWidgetItem(change+"%"))

        self.tableWidget.setItem(2,0, QtWidgets.QTableWidgetItem(str(round(secondMonthList[1],1))+" $CAD/MWh"))
        self.tableWidget.setItem(2,1, QtWidgets.QTableWidgetItem(str(round(firstMonthList[1],1))+" $CAD/MWh"))
        change = str(round(((firstMonthList[1]/secondMonthList[1])-1)*100,1))
        self.tableWidget.setItem(2,2, QtWidgets.QTableWidgetItem(change+"%"))


        self.produceSummary(firstMonthList,secondMonthList)

    def produceSummary(self,firstMonthList,secondMonthList):

        filepath = date.today().strftime('%B %d')+" Report.docx"
        if os.path.exists(filepath):
            document = Document(filepath)
        else:
            document = Document()

        if firstMonthList[1] < secondMonthList[1]:
            changeinPriceString = 'decrease'
            changeinPrice = str(round(((firstMonthList[1]/secondMonthList[1])-1)*100,1))

        elif firstMonthList[1] > secondMonthList[1]:
            changeinPriceString = 'increase'
            changeinPrice = str(round(((firstMonthList[1]/secondMonthList[1])-1)*100,1))

        if firstMonthList[2] < secondMonthList[2]:
            changeinTotalDemandString = 'decrease'
            changeinTotalDemand = str(round(((firstMonthList[2]/secondMonthList[2])-1)*100,1))
        elif firstMonthList[2] > secondMonthList[2]:
            changeinTotalDemandString = 'increase'
            changeinTotalDemand = str(round(((firstMonthList[2]/secondMonthList[2])-1)*100,1))

        if firstMonthList[3] < secondMonthList[3]:
            changeinPeakDemandString = 'decrease'
            changeinPeakDemand = str(round(((firstMonthList[3]/secondMonthList[3])-1)*100,1))
        elif firstMonthList[3] > secondMonthList[3]:
            changeinPeakDemandString = 'increase'
            changeinPeakDemand = str(round(((firstMonthList[3]/secondMonthList[3])-1)*100,1))
      

        document.add_heading("Pool Price",0)
        p = "In "+firstMonthList[0]+" the average pool price was "+ str(round(firstMonthList[1],1)) + " CAD/MWh which was a "+changeinPrice+"% "+changeinPriceString+" from the month prior. The total demand "+changeinTotalDemandString+" by "+changeinTotalDemand+"% to "+str(round(firstMonthList[2]/100,1))+" GWh and the peak demand ended up with a "+changeinPeakDemand+"% "+changeinPeakDemandString+" settling at "+ str(firstMonthList[3]) + " MWh"
        document.add_paragraph(p)
        self.summaryField.setFontPointSize(10)
        self.summaryField.setText(p)
        
        document.save(filepath)

if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_ElectricityWindow()
    ui.setupUi(MainWindow)
    MainWindow.show()
    sys.exit(app.exec_())
