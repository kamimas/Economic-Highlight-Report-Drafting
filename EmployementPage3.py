# -*- coding: utf-8 -*-

# Form implementation generated from reading ui file 'C:\Users\khossei4\Desktop\Ui Files\EmployementPage.ui'
#
# Created by: PyQt5 UI code generator 5.12.2
#
# WARNING! All changes made in this file will be lost!

from PyQt5 import QtCore, QtGui, QtWidgets

from datetime import date
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
import os.path


class Ui_MainWindow1(object):
    def setupUi(self, MainWindow1):
        MainWindow1.setObjectName("MainWindow1")
        MainWindow1.resize(995, 675)
        self.centralwidget = QtWidgets.QWidget(MainWindow1)
        self.centralwidget.setObjectName("centralwidget")
        self.dateTimeEdit = QtWidgets.QDateTimeEdit(self.centralwidget)
        self.dateTimeEdit.setGeometry(QtCore.QRect(90, 30, 194, 22))
        self.dateTimeEdit.setLayoutDirection(QtCore.Qt.LeftToRight)
        self.dateTimeEdit.setCurrentSection(
            QtWidgets.QDateTimeEdit.YearSection)
        self.dateTimeEdit.setObjectName("dateTimeEdit")
        self.dateTimeEdit_2 = QtWidgets.QDateTimeEdit(self.centralwidget)
        self.dateTimeEdit_2.setGeometry(QtCore.QRect(90, 70, 194, 22))
        self.dateTimeEdit_2.setLayoutDirection(QtCore.Qt.LeftToRight)
        self.dateTimeEdit_2.setButtonSymbols(
            QtWidgets.QAbstractSpinBox.UpDownArrows)
        self.dateTimeEdit_2.setObjectName("dateTimeEdit_2")
        self.dateTimeEdit_2.setDateTime(
            QtCore.QDateTime(date.today(), QtCore.QTime(0, 0, 0)))
        self.rbM = QtWidgets.QRadioButton(self.centralwidget)
        self.rbM.setGeometry(QtCore.QRect(340, 50, 95, 20))
        self.rbM.setChecked(True)
        self.rbM.setObjectName("rbM")
        self.rbY = QtWidgets.QRadioButton(self.centralwidget)
        self.rbY.setGeometry(QtCore.QRect(460, 50, 95, 20))
        self.rbY.setObjectName("rbY")
        self.suggestedPar = QtWidgets.QTextEdit(self.centralwidget)
        self.suggestedPar.setGeometry(QtCore.QRect(50, 380, 681, 191))
        self.suggestedPar.setObjectName("suggestedPar")
        self.table = QtWidgets.QTableWidget(self.centralwidget)
        self.table.setGeometry(QtCore.QRect(50, 130, 361, 151))
        self.table.setRowCount(3)
        self.table.setColumnCount(2)
        self.table.setObjectName("table")
        item = QtWidgets.QTableWidgetItem()
        self.table.setVerticalHeaderItem(0, item)
        item = QtWidgets.QTableWidgetItem()
        self.table.setVerticalHeaderItem(1, item)
        item = QtWidgets.QTableWidgetItem()
        self.table.setVerticalHeaderItem(2, item)
        item = QtWidgets.QTableWidgetItem()
        self.table.setHorizontalHeaderItem(0, item)
        item = QtWidgets.QTableWidgetItem()
        self.table.setHorizontalHeaderItem(1, item)
        self.table.verticalHeader().setMinimumSectionSize(33)
        self.generateBtn = QtWidgets.QPushButton(self.centralwidget)
        self.generateBtn.setGeometry(QtCore.QRect(630, 50, 93, 28))
        self.generateBtn.setObjectName("generateBtn")
        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setGeometry(QtCore.QRect(10, 30, 71, 16))
        self.label.setObjectName("label")
        self.label_2 = QtWidgets.QLabel(self.centralwidget)
        self.label_2.setGeometry(QtCore.QRect(10, 70, 55, 16))
        self.label_2.setObjectName("label_2")
        self.label_3 = QtWidgets.QLabel(self.centralwidget)
        self.label_3.setGeometry(QtCore.QRect(50, 310, 281, 31))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_3.setFont(font)
        self.label_3.setObjectName("label_3")
        self.employmentCombo = QtWidgets.QComboBox(self.centralwidget)
        self.employmentCombo.setGeometry(QtCore.QRect(480, 140, 241, 22))
        self.employmentCombo.setObjectName("employmentCombo")
        self.employmentCombo.addItem("")
        self.employmentCombo.addItem("")
        self.employmentCombo.addItem("")
        self.widget = QtWidgets.QWidget(self.centralwidget)
        self.widget.setGeometry(QtCore.QRect(790, 380, 191, 191))
        self.widget.setObjectName("widget")
        MainWindow1.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow1)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 995, 26))
        self.menubar.setObjectName("menubar")
        self.menuEmployement_Statistic_Details = QtWidgets.QMenu(self.menubar)
        self.menuEmployement_Statistic_Details.setObjectName(
            "menuEmployement_Statistic_Details")
        MainWindow1.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow1)
        self.statusbar.setObjectName("statusbar")
        MainWindow1.setStatusBar(self.statusbar)
        self.menubar.addAction(
            self.menuEmployement_Statistic_Details.menuAction())

        self.retranslateUi(MainWindow1)
        QtCore.QMetaObject.connectSlotsByName(MainWindow1)

        self.generateBtn.clicked.connect(self.generate)
        self.rbM.clicked.connect(self.switchChecked)

    def retranslateUi(self, MainWindow1):
        _translate = QtCore.QCoreApplication.translate
        MainWindow1.setWindowTitle(_translate("MainWindow1", "MainWindow"))
        self.dateTimeEdit.setDisplayFormat(
            _translate("MainWindow1", "yyyy/M/d h:mm AP"))
        self.dateTimeEdit_2.setDisplayFormat(
            _translate("MainWindow1", "yyyy/M/d h:mm AP"))
        self.rbM.setText(_translate("MainWindow1", "Monthly"))
        self.rbY.setText(_translate("MainWindow1", "Yearly"))
        item = self.table.verticalHeaderItem(0)
        item.setText(_translate("MainWindow1", "Minimum"))
        item = self.table.verticalHeaderItem(1)
        item.setText(_translate("MainWindow1", "Maximum"))
        item = self.table.verticalHeaderItem(2)
        item.setText(_translate("MainWindow1", "Total Growth "))
        item = self.table.horizontalHeaderItem(0)
        item.setText(_translate("MainWindow1", "Date"))
        item = self.table.horizontalHeaderItem(1)
        item.setText(_translate("MainWindow1", "Change"))
        self.generateBtn.setText(_translate("MainWindow1", "Generate"))
        self.label.setText(_translate("MainWindow1", "Start Date"))
        self.label_2.setText(_translate("MainWindow1", "End Date"))
        self.label_3.setText(_translate(
            "MainWindow1", "Summary of the selected timeline"))
        self.employmentCombo.setItemText(0, _translate("MainWindow1", "All"))
        self.employmentCombo.setItemText(
            1, _translate("MainWindow1", "Employment full-time"))
        self.employmentCombo.setItemText(
            2, _translate("MainWindow1", "Employment part-time"))
        self.menuEmployement_Statistic_Details.setTitle(
            _translate("MainWindow1", "Employment Statistic Details"))

    def graph(self):
        path = "Employement.xlsx"
        xl = pd.ExcelFile(path)

        # print(xl.sheet_names)

        # amCharts is opened and copied into a dataframe format named df1
        df1 = xl.parse('Data')

        province = "Alberta"

        startDate = self.dateTimeEdit.dateTime()
        startDate = startDate.toPyDateTime()

        endDate = self.dateTimeEdit_2.dateTime()
        endDate = endDate.toPyDateTime()

        Characteristic = self.employmentCombo.currentText()

        if Characteristic == 'All':
            Characteristic = 'Employment'

        # df1 is filtered to extract numbers from alberta and is stored in df2
        df2 = pd.DataFrame(data=df1, columns=[
                           'When', province, "Characteristic"])

        # irrelevant data is again filtered out and the last 5 years is taken into account
        df2 = df2[(df2['When'] >= startDate) & (df2['When'] <= endDate)
                  & (df2['Characteristic'] == Characteristic)]

        # indecies are reseted and the dataframe is saved into V1 excel sheet
        df2 = df2.reset_index(drop="True")

        plt.plot('When', province, data=df2, label=df2['Characteristic'][0])
        plt.legend()
        plt.ylabel("Employment")
        plt.title("Employment Graph")
        plt.xticks(rotation=45)
        plt.show()

    def returndataFrame(self):
        path = "Employement.xlsx"
        xl = pd.ExcelFile(path)

        # print(xl.sheet_names)

        # amCharts is opened and copied into a dataframe format named df1
        df1 = xl.parse('Data')
        df2 = pd.DataFrame(data=df1, columns=[
                           'When', 'Alberta', "Characteristic"])
        df2 = df2[(df2['Characteristic'] == 'Employment')]

        return df2

    def switchChecked(self):
        if self.rbM.isChecked() == True:
            self.rbM.isChecked() == False
            self.rbY.isChecked() == True

    def getRadio(self):
        if self.rbM.isChecked() == True:
            return 1
        elif self.rbY.isChecked() == True:
            return 2

    def monthlyAnalysis(self):
        self.table.clearContents()
        path = "Employement.xlsx"
        xl = pd.ExcelFile(path)

        # amCharts is opened and copied into a dataframe format named df1
        df1 = xl.parse('Data')

        # province, startdate, enddate and Characterisitic is used as a filter
        province = "Alberta"

        startDate = self.dateTimeEdit.dateTime()
        startDate = startDate.toPyDateTime()

        endDate = self.dateTimeEdit_2.dateTime()
        endDate = endDate.toPyDateTime()
        Characteristic = self.employmentCombo.currentText()

        if Characteristic == 'All':
            Characteristic = 'Employment'

        df2 = pd.DataFrame(data=df1, columns=[
                           'When', province, "Characteristic"])

        # irrelevant data is again filtered out and the last 5 years is taken into account
        df2 = df2[(df2['When'] >= startDate) & (df2['When'] <= endDate)
                  & (df2['Characteristic'] == Characteristic)]

        # indecies are reseted and the dataframe is saved into V1 excel sheet
        df2 = df2.reset_index(drop="True")

        secondMonthList = []
        secondMonthList.append(df2['Alberta'][0])
        secondMonthList.append(df2['When'][0])
        firstMonthList = []
        firstMonthList.append(df2['Alberta'][df2.shape[0]-1])
        firstMonthList.append(df2['When'][df2.shape[0]-1])

        ind = df2.shape[0]

        maxchangeFrom = 0
        Maximum = df2['Alberta'][0]
        MaximumDate = df2['When'][0]
        for i in range(1, (ind)):
            if df2['Alberta'][i] >= Maximum:
                maxchangeFrom = round(
                    (df2['Alberta'][i] / df2['Alberta'][i-1]), 3)

                Maximum = df2['Alberta'][i]
                MaximumDate = df2['When'][i]

        # The lowest month of the selected interval
        minchangeFrom = 0
        Minimum = df2['Alberta'][0]
        MinimumDate = df2['When'][0]
        for i in range(1, (ind)):
            if df2['Alberta'][i] <= Minimum:
                minchangeFrom = (df2['Alberta'][i] / df2['Alberta'][i-1])

                Minimum = df2['Alberta'][i]
                MinimumDate = df2['When'][i]

        # table is updated with the maximum amount
        maxchangeFrom = round((maxchangeFrom-1)*100, 1)
        maxchangeFrom = str(maxchangeFrom) + "%"

        self.table.setItem(1,   0,  QtWidgets.QTableWidgetItem(
            MaximumDate.strftime('%B') + " " + MaximumDate.strftime('%Y')))
        self.table.setItem(1,   1,  QtWidgets.QTableWidgetItem(maxchangeFrom))

        # table is updated with the minimum amount
        minchangeFrom = round((minchangeFrom-1)*100, 1)
        minchangeFrom = str(minchangeFrom) + "%"
        self.table.setItem(0,   0,  QtWidgets.QTableWidgetItem(
            MinimumDate.strftime('%B') + " " + MinimumDate.strftime('%Y')))
        self.table.setItem(0,   1,  QtWidgets.QTableWidgetItem(minchangeFrom))

        # Total growth

        TotalGrowth = df2['Alberta'][df2.shape[0]-1]/df2['Alberta'][0]
        TotalGrowth = str(round((TotalGrowth-1)*100, 1))+'%'
        self.table.setItem(2,   0,  QtWidgets.QTableWidgetItem(
            startDate.strftime('%Y' + " - "+endDate.strftime('%Y'))))
        self.table.setItem(2,   1,  QtWidgets.QTableWidgetItem(TotalGrowth))

        self.produceSummary(firstMonthList, secondMonthList, Characteristic)

    def yearlyAnalysis(self):
        self.table.clearContents()
        path = "Employement.xlsx"
        xl = pd.ExcelFile(path)

        # print(xl.sheet_names)

        # amCharts is opened and copied into a dataframe format named df1
        df1 = xl.parse('Data')

        # province, startdate, enddate and Characterisitic is used as a filter
        province = "Alberta"

        startDate = self.dateTimeEdit.dateTime()
        startDate = startDate.toPyDateTime()

        endDate = self.dateTimeEdit_2.dateTime()
        endDate = endDate.toPyDateTime()

        Characteristic = self.employmentCombo.currentText()

        if Characteristic == 'All':
            Characteristic = 'Employment'

        df2 = pd.DataFrame(data=df1, columns=[
                           'When', province, "Characteristic"])

        # irrelevant data is again filtered out and the last 5 years is taken into account
        df2 = df2[(df2['When'] >= startDate) & (df2['When'] <= endDate)
                  & (df2['Characteristic'] == Characteristic)]

        # indecies are reseted and the dataframe is saved into V1 excel sheet
        df2 = df2.reset_index(drop="True")
        ind = df2.shape[0]

        j = 0
        list1 = []
        list1.append(startDate.strftime('%Y'))

        for i in range(0, ind):
            if df2['When'][i].strftime('%Y') != list1[j]:
                j += 1
                list1.append(df2['When'][i].strftime('%Y'))

        timeYear1 = "1/1/"+list1[0]+" 12:00 AM"
        timeYear2 = "12/1/"+list1[0]+" 12:00 AM"
        df3 = df2[(df2['When'] >= (timeYear1)) & (df2['When'] <= (
            timeYear2)) & (df2['Characteristic'] == Characteristic)]
        df3 = df3.reset_index(drop="True")
        shape = df3.shape[0]-1

        change = df3['Alberta'][shape]/df3['Alberta'][0]

        # Maximum and minimum change in a given year within the interval
        Minchange = change
        Minyear = list1[0]
        Maxchange = change
        Maxyear = list1[0]

        # Total Growth from the dates specified.
        # Rate is calculated and the years mentioned
        TotalGrowthYear = list1[0] + " - " + list1[-1]
        TotalGrowth = df3['Alberta'][0]

        for i in range(1, len(list1)):
            timeYear1 = "1/1/"+list1[i]+" 12:00 AM"
            timeYear2 = "12/1/"+list1[i]+" 12:00 AM"
            df3 = df2[(df2['When'] >= (timeYear1)) & (df2['When'] <= (
                timeYear2)) & (df2['Characteristic'] == Characteristic)]
            df3 = df3.reset_index(drop="True")
            shape = df3.shape[0]-1

            if i == (len(list1)-1):
                TotalGrowth = df3['Alberta'][shape]/TotalGrowth
                TotalGrowth = round(((TotalGrowth-1)*100), 1)
                # print(TotalGrowth)
            change = df3['Alberta'][shape]/df3['Alberta'][0]

            if change >= Maxchange:
                Maxchange = change
                Maxyear = list1[i]

            if change <= Minchange:
                Minchange = change
                Minyear = list1[i]

        Maxchange = (Maxchange-1)*100

        Maxchange = str(round(Maxchange, 1)) + "%"

        self.table.setItem(1,   0,  QtWidgets.QTableWidgetItem(Maxyear))
        self.table.setItem(1,   1,  QtWidgets.QTableWidgetItem(Maxchange))

        # table is updated with the minimum amount
        Minchange = (Minchange-1)*100

        Minchange = str(round(Minchange, 1)) + "%"

        self.table.setItem(0,   0,  QtWidgets.QTableWidgetItem(Minyear))
        self.table.setItem(0,   1,  QtWidgets.QTableWidgetItem(Minchange))

        # Table is updated with the Total Growth amount
        TotalGrowth = str(TotalGrowth) + "%"
        self.table.setItem(
            2,   0,  QtWidgets.QTableWidgetItem(TotalGrowthYear))
        self.table.setItem(2,   1,  QtWidgets.QTableWidgetItem(TotalGrowth))

    def produceSummary(self, firstMonthList, secondMonthList, character):

        filepath = date.today().strftime('%B %d')+" Report.docx"
        if os.path.exists(filepath):
            document = Document(filepath)
        else:
            document = Document()

        if int(firstMonthList[0]) < int(secondMonthList[0]):
            changeinPrice = str(
                round(((firstMonthList[0]/secondMonthList[0])-1)*100, 1))
            changeinPriceString = "decrease"
        elif int(firstMonthList[0]) > int(secondMonthList[0]):
            changeinPrice = str(
                round(((firstMonthList[0]/secondMonthList[0])-1)*100, 1))
            changeinPriceString = 'increase'

        document.add_heading("Employment", 0)
        
        p = "In "+firstMonthList[1].strftime('%B %Y') +" "+character.lower()+" rate stayed at "+str(firstMonthList[0])+" which is "+changeinPrice+"% "+changeinPriceString+" from "+str(
            secondMonthList[0])+" in "+secondMonthList[1].strftime('%B %Y')
        
      
        document.add_paragraph(p)

        self.suggestedPar.setFontPointSize(10)
        self.suggestedPar.setText(p)

        document.save(filepath)

    def generate(self):

        if self.getRadio() == 1:
            self.monthlyAnalysis()
            self.graph()
        elif self.getRadio() == 2:
            self.yearlyAnalysis()


if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    MainWindow1 = QtWidgets.QMainWindow()
    ui = Ui_MainWindow1()
    ui.setupUi(MainWindow1)
    MainWindow1.show()
    sys.exit(app.exec_())
