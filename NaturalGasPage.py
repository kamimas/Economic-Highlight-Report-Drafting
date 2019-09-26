# -*- coding: utf-8 -*-

# Form implementation generated from reading ui file 'C:\Users\khossei4\Desktop\Ui Files\NaturalGasPage.ui'
#
# Created by: PyQt5 UI code generator 5.12.2
#
# WARNING! All changes made in this file will be lost!

from PyQt5                  import QtCore, QtGui, QtWidgets
from datetime               import date

import calendar
import pandas               as pd  
import matplotlib.pyplot    as plt     
import numpy                as np 
import urllib.request
import os.path
from docx import Document


class Ui_NaturalGasPage(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(918, 583)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setGeometry(QtCore.QRect(40, 30, 71, 16))
        self.label.setObjectName("label")
        self.label_2 = QtWidgets.QLabel(self.centralwidget)
        self.label_2.setGeometry(QtCore.QRect(40, 70, 55, 16))
        self.label_2.setObjectName("label_2")
        self.dateTimeEdit = QtWidgets.QDateTimeEdit(self.centralwidget)
        self.dateTimeEdit.setGeometry(QtCore.QRect(120, 30, 194, 22))
        self.dateTimeEdit.setLayoutDirection(QtCore.Qt.LeftToRight)
        self.dateTimeEdit.setCurrentSection(QtWidgets.QDateTimeEdit.YearSection)
        self.dateTimeEdit.setObjectName("dateTimeEdit")
        self.dateTimeEdit_2 = QtWidgets.QDateTimeEdit(self.centralwidget)
        self.dateTimeEdit_2.setGeometry(QtCore.QRect(120, 70, 194, 22))
        self.dateTimeEdit_2.setLayoutDirection(QtCore.Qt.LeftToRight)
        self.dateTimeEdit_2.setButtonSymbols(QtWidgets.QAbstractSpinBox.UpDownArrows)
        self.dateTimeEdit_2.setObjectName("dateTimeEdit_2")
        self.dateTimeEdit_2.setDateTime(QtCore.QDateTime(date.today(), QtCore.QTime(0, 0, 0)))


       
        self.generateBtn = QtWidgets.QPushButton(self.centralwidget)
        self.generateBtn.setGeometry(QtCore.QRect(640, 30, 93, 28))
        self.generateBtn.setObjectName("generateBtn")
        
        self.table1 = QtWidgets.QTableWidget(self.centralwidget)
        self.table1.setGeometry(QtCore.QRect(30, 170, 521, 81))
        self.table1.setShowGrid(False)
        self.table1.setGridStyle(QtCore.Qt.SolidLine)
        self.table1.setObjectName("table1")
        self.table1.setColumnCount(4)
        self.table1.setRowCount(1)
        item = QtWidgets.QTableWidgetItem()
        self.table1.setVerticalHeaderItem(0, item)
        item = QtWidgets.QTableWidgetItem()
        self.table1.setHorizontalHeaderItem(0, item)
        item = QtWidgets.QTableWidgetItem()
        self.table1.setHorizontalHeaderItem(1, item)
        item = QtWidgets.QTableWidgetItem()
        self.table1.setHorizontalHeaderItem(2, item)
        item = QtWidgets.QTableWidgetItem()
        self.table1.setHorizontalHeaderItem(3, item)
        self.table1.horizontalHeader().setVisible(True)
        self.table1.horizontalHeader().setHighlightSections(True)
        self.table1.verticalHeader().setVisible(False)
        self.label_3 = QtWidgets.QLabel(self.centralwidget)
        self.label_3.setGeometry(QtCore.QRect(30, 110, 411, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_3.setFont(font)
        self.label_3.setObjectName("label_3")
        self.suggestedPar = QtWidgets.QTextEdit(self.centralwidget)
        self.suggestedPar.setGeometry(QtCore.QRect(30, 330, 611, 191))
        self.suggestedPar.setObjectName("suggestedPar")
        self.label_4 = QtWidgets.QLabel(self.centralwidget)
        self.label_4.setGeometry(QtCore.QRect(30, 270, 281, 31))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_4.setFont(font)
        self.label_4.setObjectName("label_4")
        self.widget = QtWidgets.QWidget(self.centralwidget)
        self.widget.setGeometry(QtCore.QRect(680, 330, 191, 191))
        self.widget.setObjectName("widget")
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 918, 26))
        self.menubar.setObjectName("menubar")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)


        
        self.generateBtn.clicked.connect(self.generate)
    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
        self.label.setText(_translate("MainWindow", "Start Date"))
        self.label_2.setText(_translate("MainWindow", "End Date"))
        self.dateTimeEdit.setDisplayFormat(_translate("MainWindow", "yyyy/M/d h:mm AP"))
        self.dateTimeEdit_2.setDisplayFormat(_translate("MainWindow", "yyyy/M/d h:mm AP"))
        self.generateBtn.setText(_translate("MainWindow", "Generate"))
        item = self.table1.verticalHeaderItem(0)
        item.setText(_translate("MainWindow", "sadasdcas"))
        item = self.table1.horizontalHeaderItem(0)
        item.setText(_translate("MainWindow", "Trend"))
        item = self.table1.horizontalHeaderItem(1)
        item.setText(_translate("MainWindow", "2017"))
        item = self.table1.horizontalHeaderItem(2)
        item.setText(_translate("MainWindow", "2018"))
        item = self.table1.horizontalHeaderItem(3)
        item.setText(_translate("MainWindow", "% Change"))
        self.label_3.setText(_translate("MainWindow", "Natural Gas ($CDN/GJ)"))
        self.label_4.setText(_translate("MainWindow", "Summary of the selected timeline"))




    
    def dataFrameReport(self,startDate,endDate = date.today(),Type=2):
        

        # Data is fetched from the Natural Gas Price spreadsheet and put into df1 dataFrame variable
        path = 'NaturalGasPrice.xlsx'
        xl = pd.ExcelFile(path)
        df1 = xl.parse('Data')



        #df1 is filtered to extract numbers from alberta and is stored in df2
        df2 = pd.DataFrame(data=df1 , columns=['When','Alberta','Unit'])

        #irrelevant data is again filtered out and the last 5 years is taken into account
        df2 = df2[(df2['When'] >= startDate) & (df2['When'] <= endDate) & (df2['Unit'] =='$CDN/GJ')]

        #indecies are reseted and the dataframe is saved into V1 excel sheet
        df2 = df2.reset_index(drop="True")
        
        
        # if the user chooses the yearly option, the years are extracted
        if Type ==3 :
            return df2
        elif Type == 2:

            self.graph(df2)
            ind = df2.shape[0]
            j = 0
            list1 = []
            list1.append(startDate.strftime('%Y'))
            
            change = self.Growth(list1[j],df2)
            year = list1[j]
            for i in range(0,ind):
                if df2['When'][i].strftime('%Y') != list1[j]:
                    j += 1
                    list1.append(df2['When'][i].strftime('%Y'))
                    change1 = self.Growth(list1[j],df2)
                    if change <= change1:
                        change = change1
                        year = list1[j]

    
        year1 = int(list1[0])
        year2 = int(list1[-1])
        print(df2['When'][df2.shape[0]-1].strftime('%m'))
        self.getSpecific(str(year1),str(year2),df2['When'][0].strftime('%B'),df2['When'][df2.shape[0]-1].strftime('%B'),df2)

        

    def Growth(self,year,df2):
        timeYear1 = year+"/1/1 12:00 AM"
        timeYear2 = "12/1/"+year+" 12:00 AM"

        df3 = df2[(df2['When'] >= (timeYear1)) & (df2['When'] <= (timeYear2)) ]
        df3 = df3.reset_index(drop="True")
        

        change = df3['Alberta'][df3.shape[0]-1]/df3['Alberta'][0]
        change  = round((change-1)*100,1)
        
        return change
            

    def getSpecific(self,year1,year2,month1,month2,df2,Type=2):
        

        timeYear1 = year1+"/"+month1+"/1 12:00 AM"
        timeYear2 = year2+"/"+month2+"/1 12:00 AM"

        df3 = df2[(df2['When'] >= (timeYear1))& (df2['When'] <= timeYear2)]
        df3 = df3.reset_index(drop="True")
    
        firstYear = df3['Alberta'][0]
        secondYear = df3['Alberta'][df3.shape[0]-1]
        change = secondYear/firstYear
        if Type==2:
            self.populateTable(year1,year2,change, firstYear, secondYear,month1,month2)
        elif Type==3:
            return change

    def populateTable(self,year1, year2, change, firstYear, secondYear,month1,month2):
        #the table is populated
        
        firstYearList   = []
        secondYearList  = []

        self.table1.setHorizontalHeaderItem(1,QtWidgets.QTableWidgetItem(month1+" "+year1))
        self.table1.setHorizontalHeaderItem(2,QtWidgets.QTableWidgetItem(month2+" "+year2))

        # The lists get populated to build a paragraph
        firstYearList.append    (month1)
        secondYearList.append   (month2)


        firstYearList.append    (year1)
        secondYearList.append   (year2)


        # df1 data is inserted into the cells in the proper position
        firstYearList.append(firstYear)
        firstYear = str(firstYear)
        self.table1.setItem(0,1, QtWidgets.QTableWidgetItem(firstYear))
        


        secondYearList.append(secondYear)
        secondYear = str(secondYear)
        self.table1.setItem(0,2, QtWidgets.QTableWidgetItem(secondYear))



        change = str(round((change-1)*100,1))
        self.table1.setItem(0,3, QtWidgets.QTableWidgetItem(change+"%"))

        self.produceSummary(firstYearList,secondYearList)
        
    def generate(self):

        
        startDate = self.dateTimeEdit.dateTime()
        startDate = startDate.toPyDateTime()


        endDate = self.dateTimeEdit_2.dateTime()
        endDate = endDate.toPyDateTime()
        
        self.dataFrameReport(startDate,endDate,2)
        
    def graph(self,df1):
        plt.plot('When',"Alberta",data=df1)
        plt.xticks(rotation=45)
        plt.show()

    
    def produceSummary(self,firstMonthList,secondMonthList):
        
        filepath = date.today().strftime('%B %d')+" Report.docx"
        if os.path.exists(filepath):
            document = Document(filepath)
        else:
            document = Document()


       


        if int(firstMonthList[2]) < int(secondMonthList[2]):
            changeinPrice = str(round(((firstMonthList[2]/secondMonthList[2])-1)*100,1))
            changeinPriceString = "decrese"
        elif int(firstMonthList[2]) > int(secondMonthList[2]):
            changeinPrice = str(round(((firstMonthList[2]/secondMonthList[2])-1)*100,1))
            changeinPriceString = 'increase'
        

        document.add_heading("Natural Gas Price",0)
        p = "In "+firstMonthList[0]+" the price of natural gas stayed at "+str(firstMonthList[2])+" $CAD/GJ which is "+changeinPrice+"% "+changeinPriceString+" from "+str(secondMonthList[2])+" $CAD/GJ in "+secondMonthList[0]+" "+secondMonthList[1]
        
        document.add_paragraph(p)

        self.suggestedPar.setFontPointSize(10)
        self.suggestedPar.setText(p)
        
        document.save(filepath)

if __name__ == "__main__":
    import sys
    
    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_NaturalGasPage()
    ui.setupUi(MainWindow)
    MainWindow.show()
    sys.exit(app.exec_())
