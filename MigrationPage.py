# -*- coding: utf-8 -*-

# Form implementation generated from reading ui file 'C:\Users\khossei4\Desktop\Ui Files\MigrationPage.ui'
#
# Created by: PyQt5 UI code generator 5.12.2
#
# WARNING! All changes made in this file will be lost!

from PyQt5 import QtCore, QtGui, QtWidgets
import pandas               as pd  
import matplotlib.pyplot    as plt     
import numpy                as np 
import urllib.request
import os.path
from datetime import date
from docx import Document
class Ui_MigrationWindow(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(800, 600)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.dateTimeEdit = QtWidgets.QDateTimeEdit(self.centralwidget)
        self.dateTimeEdit.setGeometry(QtCore.QRect(130, 40, 194, 22))
        self.dateTimeEdit.setObjectName("dateTimeEdit")
        self.dateTimeEdit_2 = QtWidgets.QDateTimeEdit(self.centralwidget)
        self.dateTimeEdit_2.setGeometry(QtCore.QRect(130, 70, 194, 22))
        self.dateTimeEdit_2.setObjectName("dateTimeEdit_2")
        self.dateTimeEdit_2.setDateTime(QtCore.QDateTime(date.today(), QtCore.QTime(0, 0, 0)))

        self.label_4 = QtWidgets.QLabel(self.centralwidget)
        self.label_4.setGeometry(QtCore.QRect(20, 40, 91, 21))
        self.label_4.setObjectName("label_4")
        self.label_5 = QtWidgets.QLabel(self.centralwidget)
        self.label_5.setGeometry(QtCore.QRect(20, 70, 71, 21))
        self.label_5.setObjectName("label_5")
        self.generateBtn = QtWidgets.QPushButton(self.centralwidget)
        self.generateBtn.setGeometry(QtCore.QRect(610, 40, 93, 22))
        self.generateBtn.setObjectName("generateBtn")
        self.table1 = QtWidgets.QTableWidget(self.centralwidget)
        self.table1.setGeometry(QtCore.QRect(30, 190, 511, 81))
        self.table1.setShowGrid(False)
        self.table1.setGridStyle(QtCore.Qt.SolidLine)
        self.table1.setObjectName("table1")
        self.table1.setColumnCount(4)
        self.table1.setRowCount(1)
        item = QtWidgets.QTableWidgetItem()
        self.table1.setVerticalHeaderItem(0, item)
        item = QtWidgets.QTableWidgetItem()
        self.table1.setHorizontalHeaderItem(0, item)
        item = QtWidgets.QTableWidgetItem()
        self.table1.setHorizontalHeaderItem(1, item)
        item = QtWidgets.QTableWidgetItem()
        self.table1.setHorizontalHeaderItem(2, item)
        item = QtWidgets.QTableWidgetItem()
        self.table1.setHorizontalHeaderItem(3, item)
        self.table1.horizontalHeader().setVisible(True)
        self.table1.horizontalHeader().setHighlightSections(True)
        self.table1.verticalHeader().setVisible(False)
        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setGeometry(QtCore.QRect(30, 130, 411, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label.setFont(font)
        self.label.setObjectName("label")
        self.summaryField = QtWidgets.QTextEdit(self.centralwidget)
        self.summaryField.setGeometry(QtCore.QRect(30, 360, 561, 151))
        self.summaryField.setObjectName("summaryField")
        self.label_3 = QtWidgets.QLabel(self.centralwidget)
        self.label_3.setGeometry(QtCore.QRect(30, 300, 191, 31))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.label_3.setFont(font)
        self.label_3.setObjectName("label_3")
        self.comboDrop = QtWidgets.QComboBox(self.centralwidget)
        self.comboDrop.setGeometry(QtCore.QRect(400, 40, 161, 22))
        self.comboDrop.setObjectName("comboDrop")
        self.comboDrop.addItem("")
        self.comboDrop.addItem("")
        self.comboDrop.addItem("")
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 800, 26))
        self.menubar.setObjectName("menubar")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

        self.generateBtn.clicked.connect(self.Analysis)
    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
        self.dateTimeEdit.setDisplayFormat(_translate("MainWindow", "yyyy/M/d h:mm AP"))
        self.dateTimeEdit_2.setDisplayFormat(_translate("MainWindow", "yyyy/M/d h:mm AP"))
        self.label_4.setText(_translate("MainWindow", "Start Date"))
        self.label_5.setText(_translate("MainWindow", "End Date"))
        self.generateBtn.setText(_translate("MainWindow", "Generate"))
        item = self.table1.verticalHeaderItem(0)
        item.setText(_translate("MainWindow", "sadasdcas"))
        item = self.table1.horizontalHeaderItem(0)
        item.setText(_translate("MainWindow", "Trend"))
        item = self.table1.horizontalHeaderItem(1)
        item.setText(_translate("MainWindow", "2018"))
        item = self.table1.horizontalHeaderItem(2)
        item.setText(_translate("MainWindow", "2019"))
        item = self.table1.horizontalHeaderItem(3)
        item.setText(_translate("MainWindow", "% Change"))
        self.label.setText(_translate("MainWindow", "Net Migration"))
        self.label_3.setText(_translate("MainWindow", "Summary"))
        self.comboDrop.setItemText(0, _translate("MainWindow", "Net Migration"))
        self.comboDrop.setItemText(1, _translate("MainWindow", "Net Interprovincial Migration"))
        self.comboDrop.setItemText(2, _translate("MainWindow", "Net Immigration"))


    
    def Analysis(self):
        startDate = self.dateTimeEdit.dateTime()
        startDate = startDate.toPyDateTime()


        endDate = self.dateTimeEdit_2.dateTime()
        endDate = endDate.toPyDateTime()

        self.populateTable(self.dataFrameReport(startDate,endDate,self.comboDrop.currentText()))
        self.graph(self.dataFrameReport(startDate,endDate,self.comboDrop.currentText()))


    def dataFrameReport(self,startDate,endDate = date.today(),Type="Net Migration"):

        path="Migration.xlsx"
        xl = pd.ExcelFile(path)

        df1 = xl.parse('Data') 

        if      Type == "Net Migration":
                Type =  "NetMigration"
        elif    Type == "Net Interprovincial Migration":
                Type =  "NetInterprovincialMigration"
        elif    Type == "Net Immigration":
                Type =  "NetInternationalMigration"


        df2 = pd.DataFrame(data=df1 , columns=['When','Alberta',"Type"])
        df2 = df2[(df2['When'] >= startDate) & (df2['When'] <= endDate) & (df2['Type']==Type)]


        df2 = df2.reset_index(drop="True")
        return df2

    def populateTable(self,df1):

        # These two lists are used to 
        firstYearList   = []
        secondYearList  = []


        date1 = df1['When'][0]
        firstYearList.append(date1.strftime('%Y'))
        firstYearList.append(date1.strftime('%B'))
        date1 = date1.strftime('%Y')


        date2 = df1['When'][df1.shape[0]-1]
        secondYearList.append(date2.strftime('%Y'))
        secondYearList.append(date2.strftime('%B'))
        date2 = date2.strftime('%Y')



        

        # Headers for table 1 is set based on the last two year in the selected interval
        self.table1.setHorizontalHeaderItem(1,QtWidgets.QTableWidgetItem(date1))
        self.table1.setHorizontalHeaderItem(2,QtWidgets.QTableWidgetItem(date2))


        # df1 data is inserted into the cells in the proper position
        value1 = df1['Alberta'][0] 
        value1 = str(value1)
        self.table1.setItem(0,1, QtWidgets.QTableWidgetItem(value1))
        
        firstYearList.append(value1)



        value1 = df1['Alberta'][df1.shape[0]-1]
        value1 = str(value1)
        self.table1.setItem(0,2, QtWidgets.QTableWidgetItem(value1))

        secondYearList.append(value1)


        value1 = df1['Alberta'][df1.shape[0]-1] / df1['Alberta'][0] 
        value1 = str(round((value1 - 1) * 100,1))
        self.table1.setItem(0,3, QtWidgets.QTableWidgetItem(value1+"%"))

        self.produceSummary(firstYearList,secondYearList)
    def changeinMigration(self,df1):

        value1 = df1['Alberta'][0]
        value2 = df1['Alberta'][df1.shape[0]-1]

        change = value2/value1

        return change
    
    def produceSummary(self,firstMonthList,secondMonthList):
        
        print(int(firstMonthList[2]))
        filepath = date.today().strftime('%B %d')+" Report.docx"
        if os.path.exists(filepath):
            document = Document(filepath)
        else:
            document = Document()

        changeinPrice = str(round((int(secondMonthList[2])/int(firstMonthList[2]) - 1)*100,1))
        if int(firstMonthList[2]) > int(secondMonthList[2]):
            changeinPriceString = "decrese"
        elif int(firstMonthList[2]) < int(secondMonthList[2]):
            changeinPriceString = 'increase'
      
        print(changeinPrice)
        document.add_heading("Migration Data",0)
        p = "In "+secondMonthList[1]+" "+secondMonthList[0]+" the net migration was "+ str(secondMonthList[2]) + " which was a "+changeinPrice+"% "+changeinPriceString+" from "+firstMonthList[1]+" "+firstMonthList[0]+". "
        document.add_paragraph(p)
        self.summaryField.setFontPointSize(10)
        self.summaryField.setText(p)
        
        document.save(filepath)

    def graph(self,df1):
        
        plt.plot('When',"Alberta",data=df1,label=df1['Type'][0])
        plt.legend()
        plt.xticks(rotation=45)
        plt.ylabel("Net Population")
        plt.title("Population Graph")
        plt.show()

if __name__ == "__main__": 
    import sys
    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_MigrationWindow()
    ui.setupUi(MainWindow)
    MainWindow.show()
    sys.exit(app.exec_())
