# -*- coding: utf-8 -*-

# Form implementation generated from reading ui file 'C:\Users\khossei4\Desktop\Ui Files\GDP.ui'
#
# Created by: PyQt5 UI code generator 5.12.2
#
# WARNING! All changes made in this file will be lost!

from PyQt5 import QtCore, QtGui, QtWidgets
from datetime import date
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
import os.path
from docx import Document


class Ui_GDPWindow(object):
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(846, 642)
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.table1 = QtWidgets.QTableWidget(self.centralwidget)
        self.table1.setGeometry(QtCore.QRect(40, 190, 511, 71))
        self.table1.setShowGrid(False)
        self.table1.setGridStyle(QtCore.Qt.SolidLine)
        self.table1.setObjectName("table1")
        self.table1.setColumnCount(4)
        self.table1.setRowCount(1)
        self.table1.verticalHeader().setVisible(False)
        item = QtWidgets.QTableWidgetItem()
        self.table1.setHorizontalHeaderItem(0, item)
        item = QtWidgets.QTableWidgetItem()
        self.table1.setHorizontalHeaderItem(1, item)
        item = QtWidgets.QTableWidgetItem()
        self.table1.setHorizontalHeaderItem(2, item)
        item = QtWidgets.QTableWidgetItem()
        self.table1.setHorizontalHeaderItem(3, item)
        self.table1.horizontalHeader().setVisible(True)
        self.table1.horizontalHeader().setHighlightSections(True)
        self.label = QtWidgets.QLabel(self.centralwidget)
        self.label.setGeometry(QtCore.QRect(40, 130, 411, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label.setFont(font)
        self.label.setObjectName("label")
        self.table2 = QtWidgets.QTableWidget(self.centralwidget)
        self.table2.setGeometry(QtCore.QRect(40, 330, 511, 71))
        self.table2.setShowGrid(False)
        self.table2.setGridStyle(QtCore.Qt.SolidLine)
        self.table2.setObjectName("table2")
        self.table2.setColumnCount(4)
        self.table2.setRowCount(1)
        self.table2.verticalHeader().setVisible(False)
        item = QtWidgets.QTableWidgetItem()
        self.table2.setHorizontalHeaderItem(0, item)
        item = QtWidgets.QTableWidgetItem()
        self.table2.setHorizontalHeaderItem(1, item)
        item = QtWidgets.QTableWidgetItem()
        self.table2.setHorizontalHeaderItem(2, item)
        item = QtWidgets.QTableWidgetItem()
        self.table2.setHorizontalHeaderItem(3, item)
        self.table2.horizontalHeader().setVisible(True)
        self.table2.horizontalHeader().setHighlightSections(True)
        self.label_2 = QtWidgets.QLabel(self.centralwidget)
        self.label_2.setGeometry(QtCore.QRect(40, 270, 411, 41))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.label_2.setFont(font)
        self.label_2.setObjectName("label_2")
        self.summaryField = QtWidgets.QTextEdit(self.centralwidget)
        self.summaryField.setGeometry(QtCore.QRect(40, 460, 681, 121))
        self.summaryField.setObjectName("summaryField")
        self.label_3 = QtWidgets.QLabel(self.centralwidget)
        self.label_3.setGeometry(QtCore.QRect(50, 410, 191, 31))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.label_3.setFont(font)
        self.label_3.setObjectName("label_3")
        self.generateBtn = QtWidgets.QPushButton(self.centralwidget)
        self.generateBtn.setGeometry(QtCore.QRect(500, 50, 93, 28))
        self.generateBtn.setObjectName("generateBtn")
        self.dtEdit1 = QtWidgets.QDateTimeEdit(self.centralwidget)
        self.dtEdit1.setGeometry(QtCore.QRect(170, 40, 194, 22))
        self.dtEdit1.setObjectName("dtEdit1")
        self.dtEdit2 = QtWidgets.QDateTimeEdit(self.centralwidget)
        self.dtEdit2.setGeometry(QtCore.QRect(170, 80, 194, 22))
        self.dtEdit2.setObjectName("dtEdit2")
        self.dtEdit2.setDateTime(QtCore.QDateTime(
            date.today(), QtCore.QTime(0, 0, 0)))
        self.label_4 = QtWidgets.QLabel(self.centralwidget)
        self.label_4.setGeometry(QtCore.QRect(50, 40, 91, 21))
        self.label_4.setObjectName("label_4")
        self.label_5 = QtWidgets.QLabel(self.centralwidget)
        self.label_5.setGeometry(QtCore.QRect(50, 80, 71, 21))
        self.label_5.setObjectName("label_5")
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 846, 26))
        self.menubar.setObjectName("menubar")
        self.menuGDP_Details = QtWidgets.QMenu(self.menubar)
        self.menuGDP_Details.setObjectName("menuGDP_Details")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)
        self.menubar.addAction(self.menuGDP_Details.menuAction())

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

        self.generateBtn.clicked.connect(self.Analysis)

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "MainWindow"))
        item = self.table1.horizontalHeaderItem(0)
        item.setText(_translate("MainWindow", "Trend"))
        item = self.table1.horizontalHeaderItem(1)
        item.setText(_translate("MainWindow", "2017"))
        item = self.table1.horizontalHeaderItem(2)
        item.setText(_translate("MainWindow", "2018"))
        item = self.table1.horizontalHeaderItem(3)
        item.setText(_translate("MainWindow", "% Change"))
        self.label.setText(_translate(
            "MainWindow", "Gross Domestic Product at Basic Prices ($ billion)"))
        item = self.table2.horizontalHeaderItem(0)
        item.setText(_translate("MainWindow", "Trend"))
        item = self.table2.horizontalHeaderItem(1)
        item.setText(_translate("MainWindow", "2017"))
        item = self.table2.horizontalHeaderItem(2)
        item.setText(_translate("MainWindow", "2018"))
        item = self.table2.horizontalHeaderItem(3)
        item.setText(_translate("MainWindow", "% Change"))
        self.label_2.setText(_translate(
            "MainWindow", "Gross Domestic Product at Market Prices ($ billion)"))
        self.label_3.setText(_translate("MainWindow", "Summary"))
        self.generateBtn.setText(_translate("MainWindow", "Generate"))
        self.dtEdit1.setDisplayFormat(
            _translate("MainWindow", "yyyy/M/d h:mm AP"))
        self.dtEdit2.setDisplayFormat(
            _translate("MainWindow", "yyyy/M/d h:mm AP"))
        self.label_4.setText(_translate("MainWindow", "Start Date"))
        self.label_5.setText(_translate("MainWindow", "End Date"))
        self.menuGDP_Details.setTitle(_translate("MainWindow", "GDP Details"))

    def returnDataframe(self):
        path = "GDP.xlsx"
        xl = pd.ExcelFile(path)

        # print(xl.sheet_names)

        # amCharts is opened and copied into a dataframe format named df1
        df1 = xl.parse('Data')

        df2 = pd.DataFrame(data=df1, columns=['When', "Industries", 'Alberta'])

        df2 = df2[(df2['Industries'] == "All industries")]
        df2 = df2.reset_index(drop="True")
        year1 = df2['When'][df2.shape[0]-1].strftime('%Y')
        year2 = str(int(df2['When'][df2.shape[0]-1].strftime('%Y'))-1)

        df2 = df2[(df2['When'] >= year2) & (df2['When'] <= year1)]
        df2 = df2.reset_index(drop="True")
        return df2

    def Analysis(self):
        self.table1.clearContents()
        self.table2.clearContents()
        path = "GDP.xlsx"
        xl = pd.ExcelFile(path)

        # print(xl.sheet_names)

        # amCharts is opened and copied into a dataframe format named df1
        df1 = xl.parse('Data')

        # province, startdate, enddate and Characterisitic is used as a filter
        province = "Alberta"

        startDate = self.dtEdit1.dateTime()
        startDate = startDate.toPyDateTime()

        endDate = self.dtEdit2.dateTime()
        endDate = endDate.toPyDateTime()
        Industries = "Total gross domestic product"

        # df2 is defined and populated with dataset corresponding to Total gross domestic product at Market Price
        df2 = pd.DataFrame(data=df1, columns=['When', "Industries", province])
        df2 = df2[(df2['When'] >= startDate) & (df2['When'] <=
                                                endDate) & (df2['Industries'] == Industries)]
        df2 = df2.reset_index(drop="True")
    
        date1 = df2['When'][df2.shape[0]-2]
        date1 = date1.strftime('%Y')

        date2 = df2['When'][df2.shape[0]-1]
        date2 = date2.strftime('%Y')

        # df2 data is then inserted into table2 header for the proper dates taken from the dateTime viewer
        self.table2.setHorizontalHeaderItem(
            1, QtWidgets.QTableWidgetItem(date1))
        self.table2.setHorizontalHeaderItem(
            2, QtWidgets.QTableWidgetItem(date2))

        firstMonthList = []
        secondMonthList = []

        firstMonthList.append(df2['Alberta'][df2.shape[0]-1])
        firstMonthList.append(df2['When'][df2.shape[0]-1])

        secondMonthList.append(df2['Alberta'][0])
        secondMonthList.append(df2['When'][0])

        # df2 data is inserted into the cells in the proper position
        value1 = df2['Alberta'][df2.shape[0]-2] / 1000000000
        value1 = str(value1)
        self.table2.setItem(0, 1, QtWidgets.QTableWidgetItem(value1))

        value1 = df2['Alberta'][df2.shape[0]-1] / 1000000000
        value1 = str(value1)
        self.table2.setItem(0, 2, QtWidgets.QTableWidgetItem(value1))

        value1 = df2['Alberta'][df2.shape[0]-1] / \
            df2['Alberta'][df2.shape[0]-2]
        value1 = str(round((value1 - 1) * 100, 1))
        self.table2.setItem(0, 3, QtWidgets.QTableWidgetItem(value1+"%"))

        # df3 is defined and populated with dataset corresponding to Total gross domestic product at Basic Price
        Industries = "All industries"
        df3 = pd.DataFrame(data=df1, columns=['When', "Industries", province])
        df3 = df3[(df3['When'] >= startDate) & (df3['When'] <= endDate)
                  & (df3['Industries'] == "All industries")]

        df3 = df3.reset_index(drop="True")

        date1 = df3['When'][df3.shape[0]-2]
        date1 = date1.strftime('%Y')

        date2 = df3['When'][df3.shape[0]-1]
        date2 = date2.strftime('%Y')

        # df3 data is then inserted into table1 header for the proper dates taken from the dateTime viewer
        self.table1.setHorizontalHeaderItem(
            1, QtWidgets.QTableWidgetItem(date1))
        self.table1.setHorizontalHeaderItem(
            2, QtWidgets.QTableWidgetItem(date2))

        # df3 data is inserted into the cells in the proper position
        value1 = df3['Alberta'][df3.shape[0]-2] / 1000000000
        value1 = str(value1)
        self.table1.setItem(0, 1, QtWidgets.QTableWidgetItem(value1))

        value1 = df3['Alberta'][df3.shape[0]-1] / 1000000000
        value1 = str(value1)
        self.table1.setItem(0, 2, QtWidgets.QTableWidgetItem(value1))

        value1 = df3['Alberta'][df3.shape[0]-1] / \
            df3['Alberta'][df3.shape[0]-2]
        value1 = str(round((value1 - 1) * 100, 1))
        self.table1.setItem(0, 3, QtWidgets.QTableWidgetItem(value1+"%"))
        self.produceSummary(firstMonthList, secondMonthList)
        self.Graph(df3)

    def Graph(self, df1):
        plt.plot('When', "Alberta", data=df1)
        plt.ylabel("Gross Domestic Product($billion)")
        plt.xticks(rotation=45)
        plt.title("GDP at Market Price Graph")
        plt.show()

    def produceSummary(self, firstMonthList, secondMonthList):

        
        filepath = date.today().strftime('%B %d')+" Report.docx"
        if os.path.exists(filepath):
            document = Document(filepath)
        else:
            document = Document()

        if int(firstMonthList[0]) < int(secondMonthList[0]):
            changeinPrice = str(
                round(((firstMonthList[0]/secondMonthList[0])-1)*100, 1))
            changeinPriceString = "decrese"
        elif int(firstMonthList[0]) > int(secondMonthList[0]):
            changeinPrice = str(
                round(((firstMonthList[0]/secondMonthList[0])-1)*100, 1))
            changeinPriceString = 'increase'

        document.add_heading("GDP", 0)
        p = "In "+firstMonthList[1].strftime("%B %Y") + " the GDP stayed at "+str(firstMonthList[0])+" $CAD which is " + \
            changeinPrice+"% "+changeinPriceString+" from " + \
            str(secondMonthList[0])+" $CAD in "+secondMonthList[1].strftime("%B %Y")

        
        document.add_paragraph(p)

        self.summaryField.setFontPointSize(10)
        self.summaryField.setText(p)

        document.save(filepath)


if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    MainWindow = QtWidgets.QMainWindow()
    ui = Ui_GDPWindow()
    ui.setupUi(MainWindow)
    MainWindow.show()
    sys.exit(app.exec_())
